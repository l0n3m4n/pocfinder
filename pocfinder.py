#!/usr/bin/python3

############################
# Author: l0n3m4n          #
# Desciption: PoC Finder   #
# Version: v1.3            #
############################



import argparse
import requests
import json
import csv
import os
from docx import Document
from colorama import init, Fore, Style
from alive_progress import alive_bar
from tabulate import tabulate

#import pandas as pd
#from bs4 import BeautifulSoup


init(autoreset=True)

def fetch_exploit_info(cve):
    """fetch exploit information from an API based on CVE, returns a dictionary with description and entries."""

    api_url = f"https://api.exploit.observer/?keyword={cve}"
    try:
        with alive_bar(1, bar='smooth', spinner='dots_waves', title=Fore.LIGHTGREEN_EX + 'Fetching Data') as bar:
            response = requests.get(api_url)
            response.raise_for_status() 
            bar()
            return response.json()
    except requests.exceptions.RequestException as e:
        print(f"{Fore.RED}Error making GET request: {e}{Style.RESET_ALL}")
        return None
    except json.JSONDecodeError as e:
        print(f"{Fore.RED}Error decoding JSON response: {e}{Style.RESET_ALL}")
        return None

def save_to_json(exploit_info, filename):
    """Save exploit information to a JSON file."""
    try:
        with open(filename, 'w') as json_file:
            json.dump(exploit_info, json_file, indent=4)
        
        total_entries = sum(len(urls) for urls in exploit_info['entries'].values())
        print(f"{Fore.GREEN}Data saved to {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{filename}{Style.RESET_ALL}\n{Fore.LIGHTGREEN_EX}Total Entries: {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{total_entries}{Style.RESET_ALL}")

        
    except Exception as e:
        print(f"{Fore.RED}Error saving to {filename}: {e}{Style.RESET_ALL}")


def save_to_csv(exploit_info, filename):
    """Save exploit information to a CSV file."""
    try:
        with open(filename, 'w', newline='') as csv_file:
            writer = csv.writer(csv_file)
            writer.writerow(['File Type', 'URLs'])
            total_entries = 0
            
            for file_type, urls in exploit_info['entries'].items():
                for url in urls:
                    writer.writerow([file_type, url])
                    total_entries += 1  #
        print(f"{Fore.LIGHTGREEN_EX}Data saved to {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{filename}{Style.RESET_ALL}\n{Fore.LIGHTGREEN_EX}Total Entries: {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{total_entries}{Style.RESET_ALL}")
    except Exception as e:
        print(f"{Fore.RED}Error saving to {filename}: {e}{Style.RESET_ALL}")


def save_to_txt(exploit_info, filename):
    """Save exploit information to a TXT file."""
    try:
        with open(filename, 'w') as txt_file:
            txt_file.write(f"Description: {exploit_info['description']}\n")
            txt_file.write("Entries:\n")
            total_entries = 0
            
            for file_type, urls in exploit_info['entries'].items():
                txt_file.write(f"- {file_type}:\n")
                for url in urls:
                    txt_file.write(f"  - {url}\n")
                    total_entries += 1 
        print(f"{Fore.LIGHTGREEN_EX}Data saved to {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{filename}{Style.RESET_ALL}\n{Fore.LIGHTGREEN_EX}Total Entries: {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{total_entries}{Style.RESET_ALL}")
    except Exception as e:
        print(f"{Fore.RED}Error saving to {filename}: {e}{Style.RESET_ALL}")



def save_to_html(exploit_info, filename):
    """Save into cyberpunk looking HTML report."""
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(current_dir, 'templates', 'report.html')
        total_entries = 0  
        
        with open(template_path, 'r', encoding='utf-8') as template_file:
            template = template_file.read()

            template = template.replace('{{ description }}', exploit_info['description'])

            entries_html = ''
            for file_type, urls in exploit_info['entries'].items():
                for url in urls:
                    entries_html += f"<tr><td>{file_type}</td><td><a href='{url}' target='_blank'>{url}</a></td></tr>"
                    total_entries += 1 

            template = template.replace('{{ entries }}', entries_html)
          
        with open(filename, 'w', encoding='utf-8') as html_file:
            html_file.write(template)
        print(f"{Fore.LIGHTGREEN_EX}Data saved to {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{filename}{Style.RESET_ALL}\n{Fore.LIGHTGREEN_EX}Total Entries: {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{total_entries}{Style.RESET_ALL}")
    except FileNotFoundError as fnf_error:
        print(f"{Fore.RED}Error: File not found - {fnf_error}{Style.RESET_ALL}")
    except IOError as io_error:
        print(f"{Fore.RED}Error: IO error - {io_error}{Style.RESET_ALL}")
    except Exception as e:
        print(f"{Fore.RED}Error saving to {filename}: {e}{Style.RESET_ALL}")


def save_to_docx(exploit_info, filename):
    """Save exploit information to a DOCX file."""
    try:
        doc = Document()
        doc.add_heading('Description', level=2)
        doc.add_paragraph(exploit_info['description'])
        doc.add_heading('Entries', level=2)

        total_entries = 0
        for file_type, urls in exploit_info['entries'].items():
            doc.add_heading(file_type, level=3)
            for url in urls:
                doc.add_paragraph(url)
                total_entries += 1
        doc.save(filename)
        print(f"{Fore.LIGHTGREEN_EX}Data saved to {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{filename}{Style.RESET_ALL}\n{Fore.LIGHTGREEN_EX}Total Entries: {Style.RESET_ALL}{Fore.LIGHTBLUE_EX}{total_entries}{Style.RESET_ALL}")
    except Exception as e:
        print(f"{Fore.RED}Error saving to {filename}: {e}{Style.RESET_ALL}")



def display_results_table(exploit_info):
    """Function to display exploit information in a tabular format."""
    try:
        table_data = []
        count = 0
        
        for file_type, urls in exploit_info['entries'].items():
            for url in urls:
                if 'github.com' in url:
                    table_data.append([Fore.LIGHTCYAN_EX + file_type + Style.RESET_ALL, url])
                    count += 1
                    if count >= 20:
                        break
            if count >= 20:
                break
        
        headers = [Fore.LIGHTBLUE_EX + 'File Type', 'URLs' + Style.RESET_ALL]
        if table_data:
            colored_table = []
            for row in table_data:
                colored_row = [Fore.GREEN + str(item) + Style.RESET_ALL if 'github.com' in str(item) else str(item) for item in row]
                colored_table.append(colored_row)
            
            table = tabulate(colored_table, headers=headers, tablefmt='grid')
            print('\n' + table + '\n')
        else:
            print(f"{Fore.YELLOW}No GitHub URLs found in the exploit information.{Style.RESET_ALL}")
            
    except Exception as e:
        print(f"{Fore.RED}Error displaying results: {e}{Style.RESET_ALL}")

def main():
    banner = f"""{Fore.LIGHTGREEN_EX}
██████╗  ██████╗  ██████╗███████╗██╗███╗   ██╗██████╗ ███████╗██████╗ 
██╔══██╗██╔═══██╗██╔════╝██╔════╝██║████╗  ██║██╔══██╗██╔════╝██╔══██╗
██████╔╝██║   ██║██║     █████╗  ██║██╔██╗ ██║██║  ██║█████╗  ██████╔╝
██╔═══╝ ██║   ██║██║     ██╔══╝  ██║██║╚██╗██║██║  ██║██╔══╝  ██╔══██╗
██║     ╚██████╔╝╚██████╗██║     ██║██║ ╚████║██████╔╝███████╗██║  ██║
╚═╝      ╚═════╝  ╚═════╝╚═╝     ╚═╝╚═╝  ╚═══╝╚═════╝ ╚══════╝╚═╝  ╚═╝
             Author: l0n3m4n | PoC Finder | Version: v1.2     
{Style.RESET_ALL}"""
    parser = argparse.ArgumentParser(description='Proof of Concept Finder (PoC) and CVE details',
                                     epilog=f"{Fore.LIGHTGREEN_EX}Example usage: python3 pocfinder.py  CVE-YYYY-NNNN --export html --save CVE-2024-XXXX.html{Style.RESET_ALL}")
    print(banner)
    parser.add_argument('cve', type=str, help='CVE keyword to search for')
    parser.add_argument('-e','--export', choices=['json', 'csv', 'txt', 'html', 'docx'], required=True, metavar='', default='html', help='Export format: json, csv, txt, html, docx (default: html)')
    parser.add_argument('-s', '--save', type=str, required=True, metavar='', help='Filename to save the exported data')
    args = parser.parse_args()

    print(f"{Fore.CYAN}Fetching CVE information for CVE ID: {args.cve}...{Style.RESET_ALL}\n")
    exploit_info = fetch_exploit_info(args.cve)

    try:
        if exploit_info:
            display_results_table(exploit_info)
            if args.export == 'json':
                save_to_json(exploit_info, args.save)
            elif args.export == 'csv':
                save_to_csv(exploit_info, args.save)
            elif args.export == 'txt':
                save_to_txt(exploit_info, args.save)
            elif args.export == 'html':
                save_to_html(exploit_info, args.save)
            elif args.export == 'docx':
                save_to_docx(exploit_info, args.save)
            else:
                raise ValueError(f"{Fore.RED}Usupported export format'{args.export}'")
        else:
            raise ValueError("No exploit information available")
    except ValueError as ve:
        print(f"{Fore.RED}ValueError: {ve}{Style.RESET_ALL}")
    except Exception as e:
        print(f"{Fore.RED}An unxepected error occured: {e}{Style.RESET_ALL}")
       

if __name__ == "__main__":
    main()
